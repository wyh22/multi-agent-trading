from pathlib import Path

from tradingagents.agents.schemas import AuditIssue, AuditResult
from tradingagents.agents.utils.tool_registry import build_local_tool_groups
from tradingagents.capabilities.registry import CapabilityRegistry, CapabilitySpec
from tradingagents.conversation.store import ConversationStore
from tradingagents.orchestration.supervisor import ConversationSupervisor
from tradingagents.rag.loaders import load_documents
from tradingagents.skills.registry import BUILTIN_SKILLS


class NoStructuredLLM:
    def with_structured_output(self, _schema):
        raise NotImplementedError


def test_supervisor_falls_back_to_existing_router_when_structured_output_unavailable():
    registry = CapabilityRegistry()
    registry.register(
        CapabilitySpec(
            name="deep_stock_research",
            kind="skill",
            description="deep research",
            requires_ticker=True,
        )
    )
    supervisor = ConversationSupervisor(NoStructuredLLM(), registry)
    action = supervisor.decide(
        "请深度分析600519.SH",
        current_ticker="600519.SH",
        as_of_date="2026-09-05",
        history=[],
    )
    assert action.action == "run_deep_research"
    assert action.target == "deep_stock_research"


def test_shared_rag_tool_is_available_to_news_fundamentals_and_supervisor():
    groups = build_local_tool_groups({"rag_enabled": True})
    for group in ("news", "fundamentals", "knowledge"):
        names = [tool.name for tool in groups[group]]
        assert "search_company_knowledge" in names


def test_skill_registry_loads_declarative_manifests():
    assert {
        "deep_stock_research",
        "sector_discovery",
        "document_evidence_analysis",
        "company_comparison",
    }.issubset(BUILTIN_SKILLS)
    assert BUILTIN_SKILLS["deep_stock_research"].requires_audit is True
    assert "fundamentals" in BUILTIN_SKILLS["document_evidence_analysis"].allowed_agents


def test_research_versions_are_immutable_and_can_rollback(tmp_path):
    store = ConversationStore(tmp_path / "conversation.db")
    tid = store.ensure_thread(
        current_ticker="600519.SH",
        as_of_date="2026-09-05",
    )
    v1 = store.save_research_version(
        tid,
        {
            "ticker": "600519.SH",
            "as_of_date": "2026-09-05",
            "research_context": "context-v1",
            "final_trade_decision": "decision-v1",
        },
        audit_status="PASS",
    )
    v2 = store.save_research_version(
        tid,
        {
            "ticker": "600519.SH",
            "as_of_date": "2026-09-05",
            "research_context": "context-v2",
            "final_trade_decision": "decision-v2",
        },
        audit_status="PASS",
    )
    assert v2 > v1
    assert store.get_active_research_version(tid)["id"] == v2

    restored = store.rollback_research_version(tid)
    assert restored is not None
    assert restored["id"] == v1
    assert store.get_active_research_version(tid)["id"] == v1
    assert store.get_thread(tid)["research_context"] == "context-v1"
    versions = store.list_research_versions(tid)
    assert len(versions) == 2


def test_audit_issue_can_target_specialist_repair():
    result = AuditResult(
        verdict="REVISE",
        grounding_score=0.5,
        pit_score=1.0,
        consistency_score=0.8,
        unsupported_claims=["现金流改善缺少证据"],
        issues=[
            AuditIssue(
                issue_type="missing_evidence",
                repair_target="fundamentals",
                affected_claims=["现金流改善"],
                instruction="重新核验现金流量表与对应财报原文。",
            )
        ],
        revision_instructions=["补齐现金流证据"],
        summary="需要重新取证。",
    )
    assert result.issues[0].repair_target == "fundamentals"


def test_document_loaders_support_txt_pdf_and_docx(tmp_path):
    txt = tmp_path / "note.txt"
    txt.write_text("公司公告证据。", encoding="utf-8")
    txt_docs = load_documents(
        txt,
        ticker="600519.SH",
        publish_date="2026-09-01",
    )
    assert txt_docs[0].metadata["file_hash"]
    assert "公司公告" in txt_docs[0].text

    import fitz

    pdf_path = tmp_path / "annual.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Annual report risk disclosure")
    pdf.save(pdf_path)
    pdf.close()
    pdf_docs = load_documents(
        pdf_path,
        ticker="600519.SH",
        publish_date="2026-09-01",
        doc_type="annual_report",
    )
    assert pdf_docs
    assert pdf_docs[0].metadata["page"] == 1

    from docx import Document

    docx_path = tmp_path / "report.docx"
    document = Document()
    document.add_heading("风险因素", level=1)
    document.add_paragraph("渠道库存变化需要持续观察。")
    document.save(docx_path)
    docx_docs = load_documents(
        docx_path,
        ticker="600519.SH",
        publish_date="2026-09-01",
        doc_type="annual_report",
    )
    assert "渠道库存" in docx_docs[0].text


def test_service_exposes_knowledge_and_version_routes():
    source = (
        Path(__file__).resolve().parents[1] / "service" / "app.py"
    ).read_text(encoding="utf-8")
    assert '@app.post("/knowledge/upload")' in source
    assert '@app.get("/chat/{thread_id}/versions")' in source
    assert '@app.post("/chat/{thread_id}/rollback")' in source
