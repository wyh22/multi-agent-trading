from __future__ import annotations

import hashlib
from pathlib import Path

from tradingagents.rag.models import KnowledgeDocument


def _file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _pdf_documents(
    path: Path,
    *,
    ticker: str,
    publish_date: str,
    doc_type: str,
    file_hash: str,
) -> list[KnowledgeDocument]:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover - optional agent extra
        raise RuntimeError("PDF 解析需要 PyMuPDF，请安装 agent 可选依赖") from exc

    documents: list[KnowledgeDocument] = []
    with fitz.open(path) as pdf:
        for page_index, page in enumerate(pdf):
            text = page.get_text("text").strip()
            if not text:
                continue
            page_no = page_index + 1
            documents.append(
                KnowledgeDocument(
                    doc_id=f"{file_hash}:page:{page_no}",
                    ticker=ticker,
                    title=f"{path.stem} - 第 {page_no} 页",
                    text=text,
                    publish_date=publish_date,
                    source="uploaded-pdf",
                    url=str(path),
                    doc_type=doc_type,
                    metadata={
                        "file_name": path.name,
                        "file_hash": file_hash,
                        "page": page_no,
                    },
                )
            )
    if not documents:
        raise ValueError("PDF 未提取到可检索文本；扫描版 PDF 暂未启用 OCR")
    return documents


def _docx_documents(
    path: Path,
    *,
    ticker: str,
    publish_date: str,
    doc_type: str,
    file_hash: str,
) -> list[KnowledgeDocument]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - optional agent extra
        raise RuntimeError("DOCX 解析需要 python-docx，请安装 agent 可选依赖") from exc

    doc = Document(path)
    parts: list[str] = []
    heading_path: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = str(getattr(paragraph.style, "name", "") or "")
        if style.lower().startswith("heading"):
            heading_path = [text]
            parts.append(f"\n## {text}")
        else:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(values))
        if rows:
            parts.append(f"\n[TABLE {table_index}]\n" + "\n".join(rows))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("DOCX 未提取到可检索文本")
    return [
        KnowledgeDocument(
            doc_id=file_hash,
            ticker=ticker,
            title=path.stem,
            text=text,
            publish_date=publish_date,
            source="uploaded-docx",
            url=str(path),
            doc_type=doc_type,
            metadata={
                "file_name": path.name,
                "file_hash": file_hash,
                "heading_hint": heading_path[-1] if heading_path else "",
            },
        )
    ]


def _text_documents(
    path: Path,
    *,
    ticker: str,
    publish_date: str,
    doc_type: str,
    file_hash: str,
) -> list[KnowledgeDocument]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        raise ValueError("文档为空")
    return [
        KnowledgeDocument(
            doc_id=file_hash,
            ticker=ticker,
            title=path.stem,
            text=text,
            publish_date=publish_date,
            source="uploaded-text",
            url=str(path),
            doc_type=doc_type,
            metadata={"file_name": path.name, "file_hash": file_hash},
        )
    ]


def load_documents(
    path: str | Path,
    *,
    ticker: str,
    publish_date: str,
    doc_type: str = "user_document",
) -> list[KnowledgeDocument]:
    """Load PDF/DOCX/MD/TXT into normalized PIT-aware knowledge documents."""

    source = Path(path).expanduser()
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(source)
    suffix = source.suffix.lower()
    file_hash = _file_hash(source)
    kwargs = {
        "ticker": ticker,
        "publish_date": publish_date,
        "doc_type": doc_type,
        "file_hash": file_hash,
    }
    if suffix == ".pdf":
        return _pdf_documents(source, **kwargs)
    if suffix == ".docx":
        return _docx_documents(source, **kwargs)
    if suffix in {".txt", ".md", ".markdown"}:
        return _text_documents(source, **kwargs)
    raise ValueError(f"暂不支持的知识库文档格式: {suffix or '<none>'}")
